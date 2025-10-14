"""
DOCX Parser Module.

Модуль для извлечения таблиц и изображений из DOCX документов.

Функциональность:
- Открытие и парсинг .docx/.doc файлов
- Извлечение таблиц и сохранение в JSON формате
- Извлечение изображений и передача в парсер изображений для OCR
"""

import logging
import zipfile
from io import BytesIO
from pathlib import Path
from typing import List, Union

import pandas as pd
from PIL import Image

try:
    from docx import Document
    from docx.table import Table
except ImportError:
    raise ImportError("python-docx не установлен. Установите: pip install python-docx")

from src.parsers.img_parser import image_ocr
from src.utils.config import PARSING_DIR, HEADER_ANCHORS
from src.utils.df_utils import write_to_json, clean_dataframe
from src.utils.registry import register_parser

# Настройка логирования
logger = logging.getLogger(__name__)


def _extract_images_from_docx(docx_path: Path) -> List[Path]:
    """
    Извлекает все изображения из DOCX файла.
    
    Args:
        docx_path: Путь к DOCX файлу
        
    Returns:
        Список путей к извлеченным изображениям
    """
    extracted_paths = []
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # Получаем список всех файлов в архиве
            file_list = docx_zip.namelist()
            
            # Ищем изображения в папке word/media/
            image_files = [f for f in file_list if f.startswith('word/media/')]
            
            for i, img_file in enumerate(image_files):
                try:
                    # Читаем изображение
                    img_data = docx_zip.read(img_file)
                    
                    # Создаем уникальное имя файла в формате PNG
                    base_name = docx_path.stem
                    img_filename = f"{base_name}_extracted_img_{i+1}.png"
                    img_path = PARSING_DIR / img_filename
                    
                    # Открываем изображение через PIL и сохраняем в PNG
                    img = Image.open(BytesIO(img_data))
                    img.save(img_path, "PNG")
                    
                    extracted_paths.append(img_path)
                    logger.info(f"Извлечено изображение: {img_path}")
                    
                except Exception as e:
                    logger.warning(f"Ошибка при извлечении изображения {img_file}: {e}")
                    continue
            
            return extracted_paths
            
    except Exception as e:
        logger.error(f"Ошибка при извлечении изображений из {docx_path}: {e}")
        return []


def _normalize_text(text: str) -> str:
    """
    Нормализует текст для сравнения (как в excel_parser).
    
    Args:
        text: Исходный текст
        
    Returns:
        Нормализованный текст
    """
    if not isinstance(text, str):
        text = str(text)
    text = text.lower().replace("ё", "е")
    text = " ".join(text.split())
    return text


def _find_header_row(df: pd.DataFrame) -> int:
    """
    Ищет строку с заголовками из HEADER_ANCHORS (алгоритм из excel_parser).
    
    Args:
        df: DataFrame для поиска
        
    Returns:
        Индекс строки с заголовками или None если не найдено
    """
    try:
        best_idx = None
        best_score = 0

        # Проверяем первые 20 строк на наличие заголовков
        max_rows_to_check = min(20, len(df))

        for row_idx in range(max_rows_to_check):
            row = df.iloc[row_idx]

            # Собираем текст из всех ячеек строки
            row_text_parts = []
            for val in row:
                if pd.notna(val) and str(val).strip():
                    row_text_parts.append(str(val))

            if not row_text_parts:
                continue

            # Объединяем весь текст строки
            row_text = " ".join(row_text_parts)
            normalized_row = _normalize_text(row_text)

            # Считаем совпадения с HEADER_ANCHORS (как в excel_parser)
            score = 0
            for _, keys in HEADER_ANCHORS:
                if any(k in normalized_row for k in keys):
                    score += 1

            # Обновляем лучший результат
            if score > best_score:
                best_idx = row_idx
                best_score = score
                logger.debug(
                    f"Строка {row_idx}: score={score}, текст='{row_text[:50]}...'"
                )

        # Если score >= 2, считаем что нашли заголовок (как в excel_parser)
        if best_score >= 2:
            logger.info(
                f"Найдена строка заголовков на позиции {best_idx} "
                f"(score={best_score})"
            )
            return best_idx
        else:
            logger.debug(
                f"Заголовки из HEADER_ANCHORS не найдены "
                f"(лучший score={best_score})"
            )
            return None

    except Exception as e:
        logger.error(f"Ошибка при поиске заголовков: {e}")
        return None


def _extract_table_from_header(df: pd.DataFrame, header_row_idx: int) -> pd.DataFrame:
    """
    Извлекает таблицу начиная со строки заголовков (как в excel_parser).
    
    Args:
        df: Исходный DataFrame
        header_row_idx: Индекс строки с заголовками
        
    Returns:
        DataFrame с данными таблицы (включая строку заголовков)
    """
    try:
        # Берем таблицу начиная со строки заголовков (ВКЛЮЧАЯ её)
        table_df = df.iloc[header_row_idx:].copy()

        # Сбрасываем индекс
        table_df = table_df.reset_index(drop=True)

        logger.info(
            f"Извлечена таблица: {len(table_df)} строк (включая заголовок), "
            f"начиная со строки {header_row_idx}"
        )
        logger.debug(f"Первая строка (заголовок): {table_df.iloc[0].tolist()}")

        return table_df

    except Exception as e:
        logger.error(f"Ошибка при извлечении таблицы: {e}")
        return df


def _extract_table_data(table: Table) -> pd.DataFrame:
    """
    Извлекает данные из таблицы Word.
    
    Args:
        table: Объект таблицы python-docx
        
    Returns:
        DataFrame с данными таблицы
    """
    # Извлекаем данные из всех строк
    table_data = []
    
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            # Получаем текст из ячейки, убираем лишние пробелы
            cell_text = cell.text.strip().replace('\n', ' ')
            row_data.append(cell_text)
        table_data.append(row_data)
    
    if not table_data:
        return pd.DataFrame()
    
    # Создаем DataFrame
    df = pd.DataFrame(table_data, dtype=str)
    
    # Ищем строку с заголовками из HEADER_ANCHORS
    header_row_idx = _find_header_row(df)
    
    if header_row_idx is not None:
        # Если заголовки найдены, извлекаем таблицу начиная с них
        df = _extract_table_from_header(df, header_row_idx)
        logger.info(
            f"Таблица извлечена со строки {header_row_idx}, "
            f"игнорировано {header_row_idx} строк"
        )
    else:
        # Если заголовки не найдены, оставляем все данные как есть
        logger.info("Заголовки из HEADER_ANCHORS не найдены, используем все данные")
    
    # Убираем полностью пустые строки
    df = df.dropna(how='all').reset_index(drop=True)
    
    return df


def _save_tables_to_json(tables: List[pd.DataFrame], base_filename: str) -> List[Path]:
    """
    Сохраняет таблицы в JSON файлы используя write_to_json.
    
    Args:
        tables: Список DataFrame с таблицами
        base_filename: Базовое имя файла
        
    Returns:
        Список путей к сохраненным JSON файлам
    """
    saved_files = []
    
    for i, df in enumerate(tables):
        try:
            if df.empty:
                logger.warning(f"Таблица {i+1} пуста, пропускаем")
                continue
            
            # Создаем имя файла
            if len(tables) == 1:
                filename = f"{base_filename}_table.json"
            else:
                filename = f"{base_filename}_table_{i+1}.json"
            
            file_path = PARSING_DIR / filename
            
            # Очищаем DataFrame перед сохранением
            logger.debug(f"Cleaning DataFrame for table {i+1}, original shape: {df.shape}")
            df_cleaned = clean_dataframe(df, use_languagetool=False)
            logger.debug(f"DataFrame cleaned, new shape: {df_cleaned.shape}")
            
            # Сохраняем в JSON используя write_to_json с автоопределением заголовков
            success = write_to_json(
                file_path,
                df_cleaned,
                detect_headers=True,
                temp_dir=PARSING_DIR
            )
            
            if success:
                saved_files.append(file_path)
                logger.info(f"Таблица сохранена в JSON: {file_path}")
            else:
                logger.error(f"Не удалось сохранить таблицу {i+1} в JSON")
            
        except Exception as e:
            logger.error(f"Ошибка при сохранении таблицы {i+1}: {e}")
            continue
    
    return saved_files


@register_parser(".docx", ".doc")
def parse_docx(file_path: Union[str, Path]) -> List[Path]:
    """
    Парсит DOCX файл: извлекает таблицы и изображения.
    
    Args:
        file_path: Путь к DOCX файлу
        
    Returns:
        Список путей к обработанным файлам (JSON с таблицами + результаты OCR изображений)
    """
    file_path = Path(file_path)
    logger.info(f"Начинаем парсинг DOCX файла: {file_path}")
    
    if not file_path.exists():
        logger.error(f"Файл не найден: {file_path}")
        return []
    
    try:
        # 1. Открываем документ
        document = Document(file_path)
        base_filename = file_path.stem
        
        processed_files = []
        
        # 2. Извлекаем таблицы
        logger.info(f"Поиск таблиц в документе...")
        tables = []
        for i, table in enumerate(document.tables):
            logger.info(f"Обрабатываем таблицу {i+1}/{len(document.tables)}")
            df = _extract_table_data(table)
            
            if not df.empty:
                tables.append(df)
                logger.info(f"Извлечена таблица с размерами: {df.shape}")
            else:
                logger.warning(f"Таблица {i+1} пуста или не может быть обработана")
        
        # Сохраняем таблицы в JSON
        if tables:
            logger.info(f"Найдено таблиц: {len(tables)}")
            table_files = _save_tables_to_json(tables, base_filename)
            processed_files.extend(table_files)
        else:
            logger.info("Таблицы не найдены в документе")
        
        # 3. Извлекаем изображения
        logger.info("Поиск изображений в документе...")
        extracted_images = _extract_images_from_docx(file_path)
        
        if extracted_images:
            logger.info(f"Извлечено изображений: {len(extracted_images)}")
            
            # Обрабатываем каждое изображение через парсер изображений
            for img_path in extracted_images:
                try:
                    # Вызываем парсер изображений для OCR обработки (передаем Path объект)
                    img_result = image_ocr(img_path)
                    if img_result:
                        processed_files.extend(img_result)
                        logger.info(f"Изображение {img_path} обработано через OCR")
                    else:
                        logger.warning(f"OCR не дал результатов для {img_path}")
                        
                except Exception as e:
                    logger.warning(f"Ошибка при OCR обработке изображения {img_path}: {e}")
        else:
            logger.info("Изображения не найдены в документе")
        
        logger.info(f"Парсинг завершен. Обработано файлов: {len(processed_files)}")
        return processed_files
        
    except Exception as e:
        logger.error(f"Критическая ошибка при парсинге {file_path}: {e}")
        return []

